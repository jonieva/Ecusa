# -*- coding: utf-8 -*-
from member import Member
from docx import Document


import csv
from datetime import datetime

import dominate
from dominate.tags import *

ENGLISH = 0
SPANISH = 1
FIELDS = {
    "first_name": ["First name", "Nombre"],
    "surname": ["Surname", "Apellidos"],
    "company": ["Company", "Empresa"],
    "position": ["Position", "Cargo"],
    "keywords": ["Specialities", "Áreas de especialidad"],
    "email": ["Email", "Email"],
    "linkedin": ["Web site / LinkedIn", "Página web / LinkedIn"],
}

def generate_html(members_list, language_code):
    """ Generate the html document for a list of members
    :param members_list: list of members objects
    :param language_code: 0 (ENGLISH) or 1 (SPANISH)
    :return: html
    """
    doc = dominate.document(title='ECUSA experts guide')
    categories = ["Medicine", "Biology"]

    with doc:
        with header():
            s = '''.sort {
                      padding:8px 30px;
                      border-radius: 6px;
                      border:none;
                      /*display:inline-block;*/
                      color:#fff;
                      text-decoration: none;
                      background-color: #28a8e0;
                      height:30px;
                    }
                    .sort:hover {
                      text-decoration: none;
                      background-color:#1b8aba;
                    }
                    .sort:focus {
                      outline:none;
                    }
                    .sort:after {
                      /*display:inline-block;*/
                      width: 0;
                      height: 0;
                      border-left: 5px solid transparent;
                      border-right: 5px solid transparent;
                      border-bottom: 5px solid transparent;
                      content:"";
                      position: relative;
                      top:-10px;
                      right:-5px;
                    }
                    .sort.asc:after {
                      width: 0;
                      height: 0;
                      border-left: 5px solid transparent;
                      border-right: 5px solid transparent;
                      border-top: 5px solid #fff;
                      content:"";
                      position: relative;
                      top:4px;
                      right:-5px;
                    }
                    .sort.desc:after {
                      width: 0;
                      height: 0;
                      border-left: 5px solid transparent;
                      border-right: 5px solid transparent;
                      border-bottom: 5px solid #fff;
                      content:"";
                      position: relative;
                      top:-4px;
                      right:-5px;
                    }
                    .hidden {
                      display: none
                    }
                    '''
            style(s, type="text/css")

            script(type="text/javascript", src="http://listjs.com/no-cdn/list.js")
            script(type="text/javascript", src="http://code.jquery.com/jquery-1.11.3.min.js")

            s = '''
                $(function(){
                    var options = {
                        valueNames: ['''
            for key in FIELDS.iterkeys():
                s+= "'{0}',".format(key)
            s = s[:-1]
            s+= "]};"

            for category in categories:
                s+= "var {0}List = new List('div{0}', options);".format(category)

            s+= '''
                $('.search').keyup(function(e){
                    $('.searchDiv').css("display", "block")
                '''
            for category in categories:
                 s += '{0}List.search($(this).val());'.format(category)
            # Hide empty categories (first we have to show all of them in order that search works)
            s+= '''
                    $('div > table > tbody:not(:has(*))').parent().parent().css("display", "none")
                    });
                });'''
            # Create the script without enconding
            scr = script(type="text/javascript")
            scr.add_raw_string(s)

        with body():
            with div(id="users"):
                with div("Search for any field"):
                    input(id="searchInput", cls="search", placeholder="Search")

                i = 0
                for category in categories:
                    with div(id="div"+category, cls="searchDiv"):
                        p(category, cls="category")
                        with table():
                            # Header
                            with thead():
                                for key, values in FIELDS.iteritems():
                                    th(values[language_code], cls="sort").set_attribute("data-sort", key)
                            with tbody(cls="list"):
                                # Rows
                                for member in members_list:
                                    with tr():
                                        for key in FIELDS.iterkeys():
                                            s = eval("member.{0}".format(key))
                                            if isinstance(s, list):
                                                text = ", ".join(s)
                                                print("collection: ", text)
                                            else:
                                                text = str(s)
                                            text = text + str(i) + str(i)

                                            td(text, cls=key)
                        i+=1

    html = doc.render()
    #print(doc)
    return html






csv_file_path = "/Users/Jorge/Projects/Ecusa/responses-sample.csv"
members_list = Member.load_from_csv(csv_file_path)


#
html = generate_html(members_list, ENGLISH)
print(html)
# html = completeHtml(html2)

html_file_path = "/Users/Jorge/Projects/Ecusa/html-sample2.html"
with open(html_file_path, "w") as f:
    f.write(html)


# Member.generateDoc(members)
#
# path = '/Users/Jorge/Desktop/doc2.docx'
# f = open(path)
# document = Document(f)
# f.close()
#
# mainCats = ["Cat1", "Cat2"]
#
# for category in mainCats:
#     document.add_heading(category, 1)
#
#     style = "Medium Shading 1 Accent 3"
#     table = document.add_table(rows=0, cols=0, style=style)
#     table.autofit = False
#     table.add_column(2300000)
#     table.add_column(2200000)
#     table.add_column(2200000)
#     row = table.add_row()
#     row.cells[0].text = "Nombre"
#     #table.columns[0].width = 4839335
#
#     table.rows[0].cells[1].text = "Apellido1"
#     # table.columns[1].width = 200
#     table.rows[0].cells[2].text = "Company"
#     # table.columns[2].width = 150
#     for i in range(4):
#         member = members_list[0]
#         row = table.add_row()
#         row.cells[0].text = member.first_name
#         row.cells[1].text = member.surname
#         row.cells[2].text = member.company
#
#
#
# document.save(path)
#
# #document.tables[0].rows[0].cells[0].width
#
