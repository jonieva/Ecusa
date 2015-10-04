# -*- coding: utf-8 -*-

import csv
from datetime import datetime

import dominate
from dominate.tags import *

from docx import Document


# noinspection PyUnresolvedReferences
class Member(object):
    # Column indexes
    __REGISTERED_DATE_HEADER__ = "Timestamp"
    # ID_IX = 0
    __FIRST_NAME_HEADER__ = "Name"
    __SURNAME_HEADER__ = "Last name"
    __EMAIL_HEADER__ = "email"
    __COMPANY_HEADER__ = "Affiliation / Institution / Company"
    __POSITION_HEADER__ = "Job Title"
    __MAIN_CATEGORY_HEADER__ = "What is your professional speciality?"
    __KEYWORDS_HEADER__ = "Can you provide a few specific keywords to your professional activity?"
    __LINKEDIN_HEADER__ = "LinkedIn profile"

    ENGLISH = 0
    SPANISH = 1
    FIELDS = {
        "first_name": ["First name", "Nombre"],
        "surname": ["Surname", "Apellidos"],
        "company": ["Company", "Empresa"],
        "position": ["Position", "Cargo"],
        "keywords": ["Specialities", "Áreas de especialidad"],
        "email": ["Email", "Email"],
        "linkedin": ["Web site / LinkedIn", "Página web / LinkedIn"],
    }

    def __init__(self):
        self.header = {}

        self.id = 0
        self.registered_date = None
        self.first_name = ""
        self.surname = ""
        self.email = ""
        self.company = ""
        self.position = ""
        self.main_category = ""
        self.keywords = []
        self.linkedin = ""


    def __str__(self):
        return  "** Member {0}\n" \
                "Registered date: {1}\n" \
                "First name: {2}\n"\
                "Surname: {3}\n"\
                "Email: {4}\n"\
                "Company: {5}\n"\
                "Position: {6}\n"\
                "Main category: {7}\n"\
                "Keywords: {8}\n"\
                "LinkedIn: {9}" \
                .format(
                    self.id,
                    self.registered_date,
                    self.first_name,
                    self.surname,
                    self.email,
                    self.company,
                    self.position,
                    self.main_category,
                    "; ".join(self.keywords),
                    self.linkedin
                )

    @staticmethod
    def __getCategories__():
        return ["Cat1", "Cat2"]


    @staticmethod
    def load_from_csv(csv_file_path):
        """ Load a list of members from a csv file
        :param csv_file_path: path to the csv file
        :return: list of members
        """
        members = []
        with open(csv_file_path, 'rb') as f:
            reader = csv.reader(f)
            header = None
            for row in reader:
                if header is None:
                    header = Member.load_header(row)
                else:
                    member = Member()
                    member.header = header
                    member.load(row)
                    members.append(member)
                    #print(member)
        return members
    
    @staticmethod
    def load_header(header_row):
        """ Build the header dictionary that will be used to load a member's data from a row.
        For every property, it will store the CSV file header key and the corresponding int index in the file
        :param header_row: row that will contain the headers of the csv file
        :return: dictionary with Key-Colum_index
        """
        header = dict()
        header[Member.__REGISTERED_DATE_HEADER__] = header_row.index(Member.__REGISTERED_DATE_HEADER__)
        header[Member.__FIRST_NAME_HEADER__] = header_row.index(Member.__FIRST_NAME_HEADER__)
        header[Member.__SURNAME_HEADER__] = header_row.index(Member.__SURNAME_HEADER__)
        header[Member.__EMAIL_HEADER__] = header_row.index(Member.__EMAIL_HEADER__)
        header[Member.__COMPANY_HEADER__] = header_row.index(Member.__COMPANY_HEADER__)
        header[Member.__POSITION_HEADER__] = header_row.index(Member.__POSITION_HEADER__)
        header[Member.__MAIN_CATEGORY_HEADER__] = header_row.index(Member.__MAIN_CATEGORY_HEADER__)
        header[Member.__KEYWORDS_HEADER__] = header_row.index(Member.__KEYWORDS_HEADER__)
        header[Member.__LINKEDIN_HEADER__] = header_row.index(Member.__LINKEDIN_HEADER__)
        return header


    def load(self, row):
        """ Load the information of a member from a csv row
        :param row: csv row (as a list)
        """
        if self.header is None:
            raise Exception("Header not loaded")
        self.registered_date = datetime.strptime(row[self.header[self.__REGISTERED_DATE_HEADER__]], "%m/%d/%Y %H:%M:%S" )
        self.first_name = row[self.header[self.__FIRST_NAME_HEADER__]]
        self.surname = row[self.header[self.__SURNAME_HEADER__]]
        self.email = row[self.header[self.__EMAIL_HEADER__]]
        self.company = row[self.header[self.__COMPANY_HEADER__]]
        self.position = row[self.header[self.__POSITION_HEADER__]]
        self.main_category = row[self.header[self.__MAIN_CATEGORY_HEADER__]]
        keywords = row[self.header[self.__KEYWORDS_HEADER__]].split("\n")
        for t in keywords:
            t = t.strip()
            if t != "":
                self.keywords.append(t)
        self.linkedin = row[self.header[self.__LINKEDIN_HEADER__]]

    @staticmethod
    def generate_html(members_list, language_code):
        """ Generate the html document for a list of members
        :param members_list: list of members objects
        :param language_code: 0 (ENGLISH) or 1 (SPANISH)
        :return: html
        """
        doc = dominate.document(title='ECUSA experts guide')
        categories = ["Medicine", "Biology"]

        with doc:
            with header():
                s = '''.sort {
                          padding:8px 30px;
                          border-radius: 6px;
                          border:none;
                          /*display:inline-block;*/
                          color:#fff;
                          text-decoration: none;
                          background-color: #28a8e0;
                          height:30px;
                        }
                        .sort:hover {
                          text-decoration: none;
                          background-color:#1b8aba;
                        }
                        .sort:focus {
                          outline:none;
                        }
                        .sort:after {
                          /*display:inline-block;*/
                          width: 0;
                          height: 0;
                          border-left: 5px solid transparent;
                          border-right: 5px solid transparent;
                          border-bottom: 5px solid transparent;
                          content:"";
                          position: relative;
                          top:-10px;
                          right:-5px;
                        }
                        .sort.asc:after {
                          width: 0;
                          height: 0;
                          border-left: 5px solid transparent;
                          border-right: 5px solid transparent;
                          border-top: 5px solid #fff;
                          content:"";
                          position: relative;
                          top:4px;
                          right:-5px;
                        }
                        .sort.desc:after {
                          width: 0;
                          height: 0;
                          border-left: 5px solid transparent;
                          border-right: 5px solid transparent;
                          border-bottom: 5px solid #fff;
                          content:"";
                          position: relative;
                          top:-4px;
                          right:-5px;
                        }
                        .hidden {
                          display: none
                        }
                        '''
                style(s, type="text/css")

                script(type="text/javascript", src="http://listjs.com/no-cdn/list.js")
                script(type="text/javascript", src="http://code.jquery.com/jquery-1.11.3.min.js")

                s = '''
                    $(function(){
                        var options = {
                            valueNames: ['''
                for key in Member.FIELDS.iterkeys():
                    s+= "'{0}',".format(key)
                s = s[:-1]
                s+= "]};"

                for category in categories:
                    s+= "var {0}List = new List('div{0}', options);".format(category)

                s+= '''
                    $('.search').keyup(function(e){
                        $('.searchDiv').css("display", "block")
                    '''
                for category in categories:
                     s += '{0}List.search($(this).val());'.format(category)
                # Hide empty categories (first we have to show all of them in order that search works)
                s+= '''
                        $('div > table > tbody:not(:has(*))').parent().parent().css("display", "none")
                        });
                    });'''
                # Create the script without enconding
                scr = script(type="text/javascript")
                scr.add_raw_string(s)

            with body():
                with div(id="users"):
                    with div("Search for any field"):
                        input(id="searchInput", cls="search", placeholder="Search")

                    for category in categories:
                        with div(id="div"+category, cls="searchDiv"):
                            p(category, cls="category")
                            with table():
                                # Header
                                with thead():
                                    for key, values in Member.FIELDS.iteritems():
                                        th(values[language_code], cls="sort").set_attribute("data-sort", key)
                                with tbody(cls="list"):
                                    # Rows
                                    for member in members_list:
                                        with tr():
                                            for key in Member.FIELDS.iterkeys():
                                                s = eval("member.{0}".format(key))
                                                if isinstance(s, list):
                                                    text = ", ".join(s)
                                                    print("collection: ", text)
                                                else:
                                                    text = str(s)

                                                td(text, cls=key)

        html = doc.render()
        #print(doc)
        return html


    @staticmethod
    def generateDoc(members_list):
        # document.save("/Users/Jorge/Desktop/mydoc.docx")
        path = '/Users/Jorge/Desktop/doc2.docx'
        f = open(path)
        document = Document(f)
        f.close()

        mainCats = ["Cat1", "Cat2"]
        #document.add_heading(u'Guía de expertos de ECUSA', 0)

        for category in mainCats:
            document.add_heading(category, 1)

            style = "Medium Shading 1 Accent 3"
            table = document.add_table(rows=1, cols=3, style=style)
            table.rows[0].cells[0].text = "Nombre"
            table.columns[0].width = 40
            table.rows[0].cells[1].text = "Apellido1"
            table.columns[1].width = 100
            table.rows[0].cells[2].text = "Apellido2"
            table.columns[2].width = 150
            for i in range(4):
                member = members_list[0]
                row = table.add_row()
                row.cells[0].text = member.first_name
                row.cells[1].text = member.surname
                row.cells[2].text = member.company



        document.save(path)

